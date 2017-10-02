# -*- coding: utf-8 -*-


import os
import datetime

from docx import Document
from docx.shared import Cm, Pt


def py2docx(name_docx, name_py, *args):
    """
    Creating Excel files with Python and python-docx.

    Attributes:
    name_docx : Microsoft Word file name.
    name_py   : Application name.
    *args     : Data.
    """

    document = Document()

    core_properties = document.core_properties
    core_properties.author = 'Pedro Biel'
    core_properties.title = 'Wind action'
    core_properties.subject = 'according to FEM 2131/2132'

    # Margins.

    sections = document.sections

    for section in sections:

        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.8)

    # Style.

    style = document.styles['Normal']
    font = style.font
    font.name = 'Consolas'
    font.size = Pt(11)
    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(6)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = 1.5

    # Titel.

    document.add_heading('Wind action according to FEM 2131/2132', 0)
    paragraph = document.add_paragraph(name_py)
    paragraph = document.add_paragraph(datetime.datetime.today().strftime('%Y.%m.%d %H:%M:%S'))

    # Project.

    document.add_heading('Project', 1)

    paragraph = document.add_paragraph('Project \t \t \t \t : \t ')
    paragraph.add_run(args[0])

    paragraph = document.add_paragraph('Name \t \t \t \t \t : \t ')
    paragraph.add_run(args[1])

    paragraph = document.add_paragraph('Company \t \t \t \t : \t ')
    paragraph.add_run(args[2])

    paragraph = document.add_paragraph('Author \t \t \t \t : \t ')
    paragraph.add_run(args[3])

    paragraph = document.add_paragraph('Commentary \t \t \t \t : \t ')
    paragraph.add_run(args[4])

    if len(args) > 25:  # If the length of the args file is more than 25 the commentaries,
                        # field has more than one paragraph.

        extra_paragraphs = len(args) - 25

        for i in range(extra_paragraphs):
            paragraph = document.add_paragraph('\t \t \t \t \t : \t ')
            paragraph.add_run(args[5 + i])
            i += 1

    # Wind action.

    document.add_heading('Wind action', 1)

    # In service.

    document.add_heading('In service', 2)

    paragraph = document.add_paragraph('Design wind speed  \t \t : \t ')
    paragraph.add_run(args[-20])
    paragraph.add_run(' m/s \t \t \t § 2-2.2.1')

    paragraph = document.add_paragraph('\t \t \t \t \t : \t ')
    paragraph.add_run(args[-19])
    paragraph.add_run(' km/h')

    paragraph = document.add_paragraph('Aerodynamic wind pressure \t : \t ')
    paragraph.add_run(args[-18])
    paragraph.add_run(' kN/m² \t \t § 2-2.2.1')

    paragraph = document.add_paragraph('Ratio \t \t \t \t : \t ')
    paragraph.add_run(args[-17])

    # Travelling to storm anchoring.

    document.add_heading('Travelling to storm anchoring', 2)

    paragraph = document.add_paragraph('Design wind speed  \t \t : \t ')
    paragraph.add_run(args[-16])
    paragraph.add_run(' m/s')

    paragraph = document.add_paragraph('\t \t \t \t \t : \t ')
    paragraph.add_run(args[-15])
    paragraph.add_run(' km/h')

    paragraph = document.add_paragraph('Aerodynamic wind pressure \t : \t ')
    paragraph.add_run(args[-14])
    paragraph.add_run(' kN/m²')

    paragraph = document.add_paragraph('Ratio \t \t \t \t : \t ')
    paragraph.add_run(args[-13])

    # Out of service.

    # 0 to 20 m in height

    document.add_heading('Out service', 2)

    document.add_heading('0 to 20 m in height', 3)

    paragraph = document.add_paragraph('Design wind speed  \t \t : \t ')
    paragraph.add_run(args[-12])
    paragraph.add_run(' m/s \t \t \t § 2-2.3.6')

    paragraph = document.add_paragraph('\t \t \t \t \t : \t ')
    paragraph.add_run(args[-11])
    paragraph.add_run(' km/h')

    paragraph = document.add_paragraph('Aerodynamic wind pressure \t : \t ')
    paragraph.add_run(args[-10])
    paragraph.add_run(' kN/m² \t \t § 2-2.3.6')

    paragraph = document.add_paragraph('Ratio \t \t \t \t : \t ')
    paragraph.add_run(args[-9])

    # 20 to 100 m in height

    document.add_heading('20 to 100 m in height', 3)

    paragraph = document.add_paragraph('Design wind speed  \t \t : \t ')
    paragraph.add_run(args[-8])
    paragraph.add_run(' m/s \t \t \t § 2-2.3.6')

    paragraph = document.add_paragraph('\t \t \t \t \t : \t ')
    paragraph.add_run(args[-7])
    paragraph.add_run(' km/h')

    paragraph = document.add_paragraph('Aerodynamic wind pressure \t : \t ')
    paragraph.add_run(args[-6])
    paragraph.add_run(' kN/m² \t \t § 2-2.3.6')

    paragraph = document.add_paragraph('Ratio \t \t \t \t : \t ')
    paragraph.add_run(args[-5])

    # More than 100 m in height

    document.add_heading('More than 100 m in height', 3)

    paragraph = document.add_paragraph('Design wind speed  \t \t : \t ')
    paragraph.add_run(args[-4])
    paragraph.add_run(' m/s \t \t \t § 2-2.3.6')

    paragraph = document.add_paragraph('\t \t \t \t \t : \t ')
    paragraph.add_run(args[-3])
    paragraph.add_run(' km/h')

    paragraph = document.add_paragraph('Aerodynamic wind pressure \t : \t ')
    paragraph.add_run(args[-2])
    paragraph.add_run(' kN/m² \t \t § 2-2.3.6')

    paragraph = document.add_paragraph('Ratio \t \t \t \t : \t ')
    paragraph.add_run(args[-1])

    # Save document

    document.save(name_docx)

    # Starts Excel

    #os.system(name_docx)


if __name__ == '__main__':

    data = ['12345', 'Estructuras', 'TAIM WESERs', 'Dr Nos', 'No comments 1.', 'No comments 2.', 'No 3.', 'No 4.',
            '20', '72', '0,25', '1,000',
            '28', '100', '0,48', '1,960',
            '36', '129', '0,79', '3,240',
            '42', '151', '1,08', '4,410',
            '46', '165', '1,30', '5,290'
            ]

    py2docx('aaa.docx', *data)

    #os.system(r'start D:\Python\Apps\Cierzo\codigo\scripts\aaa.docx')